from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
import docx
from datetime import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
import os

# 搜索關鍵詞列表
keywords = ['拍攝', '影片', '紀錄片', '平面影像', '影音', '影片製作', 
           '素材製作', '短片', '攝影', '錄影', '錄製']

base_url = 'https://web.pcc.gov.tw/prkms/tender/common/basic/readTenderBasic'

def setup_driver():
    chrome_options = Options()
    chrome_options.add_argument('--headless')
    chrome_options.add_argument('--no-sandbox')
    chrome_options.add_argument('--disable-dev-shm-usage')
    try:
        driver = webdriver.Chrome(options=chrome_options)
        print("ChromeDriver 初始化成功")
        return driver
    except Exception as e:
        print(f"ChromeDriver 初始化失敗: {str(e)}")
        raise

def search_tender(keyword):
    driver = setup_driver()
    driver.get(base_url)
    
    try:
        wait = WebDriverWait(driver, 10)
        print(f"開始搜尋關鍵詞: {keyword}")
        
        tender_name = wait.until(EC.presence_of_element_located((By.ID, 'tenderName')))
        tender_name.send_keys(keyword)
        print("已輸入標案名稱")
        
        tender_status = wait.until(EC.presence_of_element_located((By.ID, 'tenderStatus')))
        tender_status.send_keys('等標期內')
        print("已選擇等標期內")
        
        search_button = wait.until(EC.element_to_be_clickable((By.ID, 'searchButton')))
        search_button.click()
        print("已點擊搜尋按鈕")
        
        time.sleep(3)
        
        results = driver.find_elements(By.CLASS_NAME, 'tender-result-row')
        result_list = []
        for result in results:
            try:
                title = result.find_element(By.CLASS_NAME, 'tender-title').text
                date = result.find_element(By.CLASS_NAME, 'tender-date').text
                result_list.append({'title': title, 'date': date})
            except Exception as e:
                print(f"解析結果時出錯: {str(e)}")
        print(f"找到 {len(result_list)} 個結果")
        return result_list
        
    except Exception as e:
        print(f"搜尋 {keyword} 時出錯: {str(e)}")
        return []
        
    finally:
        driver.quit()

def create_word_document(results):
    try:
        doc = docx.Document()
        doc.add_heading('標案搜尋結果', 0)
        doc.add_paragraph(f'生成日期: {datetime.now().strftime("%Y-%m-%d")}')
        
        for result in results:
            doc.add_paragraph(f"標案名稱: {result['title']}")
            doc.add_paragraph(f"公告日期: {result['date']}")
            doc.add_paragraph("-" * 50)
        
        filename = f'tender_results_{datetime.now().strftime("%Y%m%d")}.docx'
        doc.save(filename)
        print(f"已生成文件: {filename}")
        return filename
    except Exception as e:
        print(f"生成 Word 文件失敗: {str(e)}")
        raise

def send_email(filename):
    sender_email = os.environ.get('SENDER_EMAIL')
    sender_password = os.environ.get('SENDER_PASSWORD')
    receiver_email = "haap0716@gmail.com"
    
    if not sender_email or not sender_password:
        print("缺少郵件發送憑證（SENDER_EMAIL 或 SENDER_PASSWORD）")
        return
    
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = receiver_email
    msg['Subject'] = f'每日標案搜尋結果 - {datetime.now().strftime("%Y-%m-%d")}'
    
    try:
        with open(filename, "rb") as f:
            part = MIMEApplication(f.read(), Name=os.path.basename(filename))
            part['Content-Disposition'] = f'attachment; filename="{os.path.basename(filename)}"'
            msg.attach(part)
        
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(sender_email, sender_password)
            server.send_message(msg)
            print("郵件發送成功")
    except Exception as e:
        print(f"郵件發送失敗: {str(e)}")

def main():
    all_results = []
    for keyword in keywords:
        results = search_tender(keyword)
        all_results.extend(results)
    
    unique_results = []
    seen_titles = set()
    for result in all_results:
        if result['title'] not in seen_titles:
            unique_results.append(result)
            seen_titles.add(result['title'])
    print(f"總共找到 {len(unique_results)} 個唯一結果")
    
    if unique_results:
        filename = create_word_document(unique_results)
        send_email(filename)
        if os.path.exists(filename):
            os.remove(filename)
            print("臨時文件已刪除")
    else:
        print("沒有找到任何結果")

if __name__ == "__main__":
    main()
           
